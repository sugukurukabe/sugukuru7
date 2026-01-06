from docx import Document
from docx.shared import Pt, Cm
import os
import logging

logger = logging.getLogger(__name__)

class WordTemplateProcessor:
    """
    Wordテンプレート処理
    {{氏名}} などのプレースホルダーを置換します。
    """
    
    def __init__(self, template_path: str):
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        self.doc = Document(template_path)
    
    def fill_placeholders(self, data: dict):
        """
        プレースホルダーを置換（段落とテーブル両方）
        """
        # 段落内の置換
        for paragraph in self.doc.paragraphs:
            self._replace_in_element(paragraph, data)
        
        # テーブル内の置換
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_element(paragraph, data)
    
    def _replace_in_element(self, paragraph, data):
        """段落内のテキスト置換。書式を維持するため Run レベルで処理が必要な場合があるが、まずは単純化"""
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                # 書式を壊さないための工夫（全置換が望ましいが、Runが分かれている場合に備えて）
                # ここでは簡易的に paragraph.text 全体を置換（これだとRunの個別書式が一部消える可能性がある）
                # より高度な実装では Run を跨ぐプレースホルダー検出が必要
                paragraph.text = paragraph.text.replace(placeholder, str(value or ""))
    
    def save(self, output_path: str):
        """保存"""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path
